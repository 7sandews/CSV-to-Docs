import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import pandas as pd

def generate_docx(csv_path, output_path):
    try:
        df_sorted = pd.read_csv(csv_path)

        doc = Document()
        doc.add_heading("MCQ Question Bank", level=1)

        question_number = 1

        for subdomain, group in df_sorted.groupby('Sub-domain'):
            for _, row in group.iterrows():
                doc.add_heading(f"Sub Domain – {subdomain}", level=2)
                doc.add_paragraph(f"Complexity – {row['Complexity'].capitalize()}", style='Intense Quote')
                
                doc.add_paragraph(f"Q. {question_number}) {row['Question']}", style='Normal')
                doc.add_paragraph(f"A) {row['Choice1']}")
                doc.add_paragraph(f"B) {row['Choice2']}")
                doc.add_paragraph(f"C) {row['Choice3']}")
                doc.add_paragraph(f"D) {row['Choice4']}")
                
                doc.add_paragraph("")  # spacing
                question_number += 1

        doc.save(output_path)
        messagebox.showinfo("Success", f"Word document saved to:\n{output_path}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to generate document:\n{str(e)}")

def browse_csv():
    file_path = filedialog.askopenfilename(filetypes=[("CSV Files", "*.csv")])
    if file_path:
        csv_entry.delete(0, tk.END)
        csv_entry.insert(0, file_path)

def browse_output():
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if file_path:
        output_entry.delete(0, tk.END)
        output_entry.insert(0, file_path)

def run_script():
    csv_path = csv_entry.get()
    output_path = output_entry.get()

    if not csv_path or not output_path:
        messagebox.showwarning("Missing Info", "Please select both input CSV and output file path.")
        return

    generate_docx(csv_path, output_path)

# GUI setup
root = tk.Tk()
root.title("Sample Question Sheet Generator")
root.geometry("550x250")

# CSV input
tk.Label(root, text="1. Select your CSV file:").pack(pady=(10, 0))
csv_entry = tk.Entry(root, width=60)
csv_entry.pack(padx=10)
tk.Button(root, text="Browse CSV", command=browse_csv).pack(pady=2)

# Output path
tk.Label(root, text="2. Choose output .docx file location:").pack(pady=(10, 0))
output_entry = tk.Entry(root, width=60)
output_entry.pack(padx=10)
tk.Button(root, text="Save As", command=browse_output).pack(pady=2)

# Run Button
generate_button = tk.Button(root, 
    text="Generate Word Document", 
    command=run_script,
    bg="#2196F3",   # blue background
    fg="white",     # white text
    font=('Arial', 12, 'bold'),
    activebackground="#1976D2",  # darker blue when clicked
    activeforeground="white"     # stay white when clicked
)
generate_button.pack(pady=20)

root.mainloop()