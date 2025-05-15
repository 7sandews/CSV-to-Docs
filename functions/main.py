import functions_framework
from flask import request, send_file
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import io
import tempfile

@functions_framework.http
def generateDocx(request):
    # Set CORS headers for the preflight request
    if request.method == 'OPTIONS':
        # Allows GET requests from any origin with the Content-Type
        headers = {
            'Access-Control-Allow-Origin': '*',
            'Access-Control-Allow-Methods': 'POST',
            'Access-Control-Allow-Headers': 'Content-Type',
            'Access-Control-Max-Age': '3600'
        }
        return ('', 204, headers)

    # Set CORS headers for the main request
    headers = {
        'Access-Control-Allow-Origin': '*'
    }

    try:
        # Get the CSV file from the request
        csv_file = request.files['csv']
        
        # Read the CSV data
        df = pd.read_csv(csv_file)
        
        # Create a new Word document
        doc = Document()
        
        # Set up document styles
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # Add title
        title = doc.add_heading("MCQ Question Bank", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        question_number = 1
        
        # Group by subdomain and add questions
        for subdomain, group in df.groupby('Sub-domain'):
            # Add subdomain heading
            doc.add_heading(f"Sub Domain – {subdomain}", level=2)
            
            for _, row in group.iterrows():
                # Add complexity
                complexity = doc.add_paragraph(f"Complexity – {row['Complexity'].capitalize()}")
                complexity.style = 'Intense Quote'
                
                # Add question
                question = doc.add_paragraph(f"Q. {question_number}) {row['Question']}")
                question.style = 'Normal'
                
                # Add choices
                doc.add_paragraph(f"A) {row['Choice1']}")
                doc.add_paragraph(f"B) {row['Choice2']}")
                doc.add_paragraph(f"C) {row['Choice3']}")
                doc.add_paragraph(f"D) {row['Choice4']}")
                
                # Add spacing
                doc.add_paragraph()
                question_number += 1
        
        # Save the document to a BytesIO object
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return send_file(
            docx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='MCQ_Question_Bank.docx'
        )
        
    except Exception as e:
        return (str(e), 500, headers) 