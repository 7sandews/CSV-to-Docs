from flask import Flask, request, send_file
from docx import Document
import pandas as pd
import io
import tempfile
import os

app = Flask(__name__, static_folder='docs')

@app.route('/')
def index():
    return app.send_static_file('index.html')

@app.route('/generate-docx', methods=['POST'])
def generate_docx():
    try:
        # Get the CSV file from the request
        csv_file = request.files['csv']
        
        # Read the CSV data
        df = pd.read_csv(csv_file)
        
        # Create a new Word document
        doc = Document()
        doc.add_heading("MCQ Question Bank", level=1)
        
        question_number = 1
        
        # Group by subdomain and add questions
        for subdomain, group in df.groupby('Sub-domain'):
            doc.add_heading(f"Sub Domain – {subdomain}", level=2)
            
            for _, row in group.iterrows():
                doc.add_paragraph(f"Complexity – {row['Complexity'].capitalize()}", style='Intense Quote')
                doc.add_paragraph(f"Q. {question_number}) {row['Question']}", style='Normal')
                doc.add_paragraph(f"A) {row['Choice1']}")
                doc.add_paragraph(f"B) {row['Choice2']}")
                doc.add_paragraph(f"C) {row['Choice3']}")
                doc.add_paragraph(f"D) {row['Choice4']}")
                doc.add_paragraph("")  # spacing
                question_number += 1
        
        # Save the document to a BytesIO object
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return send_file(
            docx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='MCQ_Question_Bank.docx'
        )
        
    except Exception as e:
        return str(e), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000) 